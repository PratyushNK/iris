from __future__ import annotations

import base64
from io import BytesIO

from docx import Document
from docx.document import Document as DocumentType
from docx.shared import Inches, Pt

from agent.agent_state import IrisAgentState as AgentState


def _slugify_filename(title: str) -> str:
    cleaned = [character.lower() if character.isalnum() else "_" for character in title]
    slug = "".join(cleaned)
    while "__" in slug:
        slug = slug.replace("__", "_")
    slug = slug.strip("_")
    return slug or "iris_deliverable"


def _apply_document_defaults(document: DocumentType) -> None:
    styles = document.styles
    normal_style = styles["Normal"]
    # python-docx exposes font on the runtime style object even though the stub is looser here.
    normal_style.font.name = "Aptos"  # type: ignore[attr-defined]
    normal_style.font.size = Pt(11)  # type: ignore[attr-defined]

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)


def _write_sections(document: DocumentType, sections: list) -> None:
    for section in sections:
        document.add_heading(section.heading, level=1)
        for paragraph in section.paragraphs:
            document.add_paragraph(paragraph)
        for bullet in section.bullets:
            document.add_paragraph(bullet, style="List Bullet")


def generate_report_node(state: AgentState):
    title = state.get("title", "Iris Deliverable")
    sections = state.get("docx_sections", [])

    document = Document()
    _apply_document_defaults(document)

    document.core_properties.title = title
    document.core_properties.author = "Iris Autonomous Agent"

    document.add_heading(title, level=0)
    intro = document.add_paragraph()
    intro.add_run("Generated automatically by the Iris autonomous agent.").italic = True

    if sections:
        _write_sections(document, sections)
    else:
        document.add_paragraph(state.get("docx_content", ""))

    buffer = BytesIO()
    document.save(buffer)
    docx_bytes = buffer.getvalue()

    encoded_string = base64.b64encode(docx_bytes).decode("utf-8")

    return {
        "docx_file_b64": encoded_string,
        "docx_filename": f"{_slugify_filename(title)}.docx",
    }
